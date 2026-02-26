from __future__ import annotations

import io
import re
from typing import Any

import numpy as np
import pandas as pd

from app.models.schemas import RiskAssessmentRequest, RiskAssessmentResponse
from app.services.risk_model_manager import risk_model_manager


def _clamp(value: float, low: float, high: float) -> float:
    return max(low, min(high, value))


def _normalize_text(value: str) -> str:
    return value.strip().lower()


def _purpose_profile(purpose: str) -> tuple[str, float, int, int, float]:
    text = _normalize_text(purpose)
    if any(x in text for x in ["home", "house", "property"]):
        return ("Home Loan", 9.0, 120, 300, 16.0)
    if any(x in text for x in ["education", "study", "college", "course"]):
        return ("Education Loan", 10.0, 36, 120, 20.0)
    if any(x in text for x in ["business", "inventory", "shop", "working capital"]):
        return ("Business Loan", 14.0, 12, 60, 28.0)
    if any(x in text for x in ["vehicle", "bike", "car", "auto"]):
        return ("Vehicle Loan", 11.0, 24, 84, 22.0)
    if any(x in text for x in ["medical", "health", "hospital"]):
        return ("Medical Personal Loan", 15.0, 12, 48, 30.0)
    return ("Personal Loan", 16.0, 12, 60, 32.0)


def _occupation_risk(occupation: str) -> tuple[float, str]:
    text = _normalize_text(occupation)
    if any(x in text for x in ["government", "govt", "teacher", "bank employee"]):
        return (12.0, "Your occupation profile appears stable, which supports repayment reliability.")
    if any(x in text for x in ["salaried", "software", "engineer", "private employee"]):
        return (16.0, "A regular salaried income pattern generally improves repayment consistency.")
    if any(x in text for x in ["self", "business", "shop", "vendor", "trader"]):
        return (24.0, "Self-employment can involve income variability, so risk is treated as moderate.")
    if any(x in text for x in ["daily wage", "freelancer", "contract"]):
        return (34.0, "This occupation type may have uneven monthly cash flow, increasing repayment uncertainty.")
    if "student" in text:
        return (36.0, "Student profiles usually have limited independent repayment capacity at present.")
    return (26.0, "Occupation profile indicates a moderate repayment risk band.")


def _age_risk(age: int) -> tuple[float, str]:
    if 23 <= age <= 55:
        return (14.0, "Your age bracket is typically aligned with stable earning years.")
    if 18 <= age < 23:
        return (28.0, "Early-career age bands often have developing income stability.")
    if 56 <= age <= 65:
        return (26.0, "This age band can reduce lender flexibility for longer tenures.")
    return (34.0, "This profile may require shorter tenure and tighter lending terms.")


def _predict_financial_condition(occupation: str, age: int) -> tuple[float, float, str]:
    # Simple prediction model using occupation and age only.
    occupation_risk, _ = _occupation_risk(occupation)
    age_risk, _ = _age_risk(age)
    condition_risk = (0.62 * occupation_risk) + (0.38 * age_risk)
    condition_score = _clamp(100 - condition_risk, 5.0, 95.0)

    if condition_score >= 72:
        label = "Strong"
    elif condition_score >= 52:
        label = "Moderate"
    else:
        label = "Vulnerable"

    note = (
        f"Financial-condition model (occupation + age) indicates a {label.lower()} profile "
        f"with score {condition_score:.1f}/100."
    )
    return condition_score, condition_risk, note


def _cibil_risk(cibil_score: int, estimated: bool) -> tuple[float, str]:
    source = "estimated from your profile" if estimated else "provided by you"
    if cibil_score >= 780:
        return (8.0, f"The CIBIL score ({cibil_score}, {source}) is strong and supports high lender confidence.")
    if cibil_score >= 720:
        return (14.0, f"The CIBIL score ({cibil_score}, {source}) is good and supports approval probability.")
    if cibil_score >= 680:
        return (22.0, f"The CIBIL score ({cibil_score}, {source}) is fair, indicating moderate credit risk.")
    if cibil_score >= 620:
        return (34.0, f"The CIBIL score ({cibil_score}, {source}) is below ideal and may reduce approval odds.")
    return (46.0, f"The CIBIL score ({cibil_score}, {source}) is low, which materially increases risk.")


def _emi_pressure_risk(monthly_income: float, existing_emis: float) -> tuple[float, str]:
    emi_ratio = existing_emis / max(monthly_income, 1.0)
    if emi_ratio <= 0.2:
        return (10.0, "Current EMI commitments are light relative to monthly income.")
    if emi_ratio <= 0.35:
        return (20.0, "Current EMI commitments are manageable, though repayment headroom is moderate.")
    if emi_ratio <= 0.5:
        return (34.0, "Existing EMI obligations consume a significant portion of income.")
    return (48.0, "High existing EMI obligations create strong repayment pressure.")


def _loan_burden_risk(monthly_income: float, loan_amount: float) -> tuple[float, str]:
    loan_to_annual_income = loan_amount / max(monthly_income * 12.0, 1.0)
    if loan_to_annual_income <= 0.5:
        return (12.0, "Requested loan size is modest relative to estimated annual income.")
    if loan_to_annual_income <= 1.0:
        return (20.0, "Requested loan size appears reasonable for the current income profile.")
    if loan_to_annual_income <= 1.8:
        return (32.0, "Requested loan size is high relative to annual income.")
    return (44.0, "Requested loan size is very high relative to annual income.")


def _expense_risk(monthly_income: float, monthly_expenses: float) -> tuple[float, str]:
    expense_ratio = monthly_expenses / max(monthly_income, 1.0)
    if expense_ratio <= 0.45:
        return (10.0, "Monthly expense levels are well within income capacity.")
    if expense_ratio <= 0.65:
        return (20.0, "Monthly expense levels are moderate relative to income.")
    if expense_ratio <= 0.8:
        return (34.0, "Monthly expense levels are high and may pressure repayments.")
    return (48.0, "Very high expense levels leave limited room for additional EMI.")


def _savings_risk(monthly_income: float, current_savings: float) -> tuple[float, str]:
    savings_ratio = current_savings / max(monthly_income, 1.0)
    if savings_ratio >= 6:
        return (8.0, "Current savings provide a strong financial cushion for repayment continuity.")
    if savings_ratio >= 3:
        return (16.0, "Savings buffer is healthy and supports repayment resilience.")
    if savings_ratio >= 1:
        return (26.0, "Savings are limited; increasing reserves would further reduce risk.")
    return (40.0, "Low savings increase vulnerability to income and expense shocks.")


def _estimate_cibil(payload: RiskAssessmentRequest) -> int:
    income_ratio = _clamp(payload.monthly_income / 100000.0, 0.0, 1.0)
    emi_ratio = _clamp(payload.existing_emis / max(payload.monthly_income, 1.0), 0.0, 1.2)
    expense_ratio = _clamp(payload.monthly_expenses / max(payload.monthly_income, 1.0), 0.0, 1.4)
    savings_months = _clamp(payload.current_savings / max(payload.monthly_income, 1.0), 0.0, 12.0)
    loan_burden = _clamp(payload.loan_amount / max(payload.monthly_income * 12.0, 1.0), 0.0, 3.0)

    occ_risk, _ = _occupation_risk(payload.occupation)
    age_risk, _ = _age_risk(payload.age)

    score = (
        675
        + 35 * income_ratio
        - 95 * emi_ratio
        - 75 * max(expense_ratio - 0.45, 0.0)
        + 28 * (savings_months / 6.0 if savings_months <= 6 else 1.0)
        - 45 * max(loan_burden - 0.8, 0.0)
        - 0.8 * occ_risk
        - 0.5 * age_risk
    )
    return int(round(_clamp(score, 520, 790)))


def _emi(principal: float, annual_rate: float, tenure_months: int) -> float:
    monthly_rate = annual_rate / 12 / 100
    if monthly_rate == 0:
        return principal / tenure_months
    factor = (1 + monthly_rate) ** tenure_months
    return principal * monthly_rate * factor / (factor - 1)


def _pick_tenure(
    loan_amount: float,
    annual_rate: float,
    min_tenure_months: int,
    max_tenure_months: int,
    monthly_income: float,
    existing_emis: float,
    monthly_expenses: float,
) -> tuple[int, float]:
    step = 12 if max_tenure_months > 96 else 6
    tenure_options = list(range(min_tenure_months, max_tenure_months + 1, step))
    if tenure_options[-1] != max_tenure_months:
        tenure_options.append(max_tenure_months)

    affordable_emi = max(monthly_income - monthly_expenses - existing_emis, monthly_income * 0.08)
    chosen_tenure = max_tenure_months
    chosen_emi = _emi(loan_amount, annual_rate, chosen_tenure)

    for tenure in tenure_options:
        emi = _emi(loan_amount, annual_rate, tenure)
        if emi <= affordable_emi:
            chosen_tenure = tenure
            chosen_emi = emi
            break
    return chosen_tenure, chosen_emi


def assess_risk(payload: RiskAssessmentRequest) -> RiskAssessmentResponse:
    purpose_loan_type, annual_rate, min_tenure, max_tenure, purpose_risk_score = _purpose_profile(
        payload.purpose
    )

    emi_risk, emi_msg = _emi_pressure_risk(payload.monthly_income, payload.existing_emis)
    size_risk, size_msg = _loan_burden_risk(payload.monthly_income, payload.loan_amount)
    expense_risk, expense_msg = _expense_risk(payload.monthly_income, payload.monthly_expenses)
    savings_risk, savings_msg = _savings_risk(payload.monthly_income, payload.current_savings)
    condition_score, condition_risk, condition_msg = _predict_financial_condition(
        payload.occupation, payload.age
    )
    cibil_estimated = payload.cibil_score is None
    cibil_score_used = _estimate_cibil(payload) if cibil_estimated else int(payload.cibil_score)
    cibil_risk, cibil_msg = _cibil_risk(cibil_score_used, estimated=cibil_estimated)

    weighted_components: list[tuple[float, float]] = [
        (emi_risk, 0.20),
        (size_risk, 0.18),
        (expense_risk, 0.16),
        (savings_risk, 0.12),
        (condition_risk, 0.17),
        (purpose_risk_score, 0.07),
    ]
    weighted_components.append((cibil_risk, 0.10))

    total_weight = sum(w for _, w in weighted_components)
    risk_score = sum(score * weight for score, weight in weighted_components) / max(total_weight, 1e-9)

    default_probability = round(_clamp(risk_score, 3.0, 95.0), 2)
    approval_probability = round(100 - default_probability, 2)
    if default_probability < 30:
        risk_category = "Low"
    elif default_probability < 60:
        risk_category = "Medium"
    else:
        risk_category = "High"

    suggested_tenure, estimated_emi = _pick_tenure(
        loan_amount=payload.loan_amount,
        annual_rate=annual_rate,
        min_tenure_months=min_tenure,
        max_tenure_months=max_tenure,
        monthly_income=payload.monthly_income,
        existing_emis=payload.existing_emis,
        monthly_expenses=payload.monthly_expenses,
    )

    component_messages = [
        ("existing_emi", emi_risk, emi_msg),
        ("loan_size", size_risk, size_msg),
        ("monthly_expenses", expense_risk, expense_msg),
        ("current_savings", savings_risk, savings_msg),
        ("financial_condition", condition_risk, condition_msg),
        ("cibil", cibil_risk, cibil_msg),
    ]
    primary_driver = max(component_messages, key=lambda x: x[1])[2]
    emi_ratio = payload.existing_emis / max(payload.monthly_income, 1.0)
    expense_ratio = payload.monthly_expenses / max(payload.monthly_income, 1.0)
    savings_months = payload.current_savings / max(payload.monthly_income, 1.0)
    loan_to_annual_income = payload.loan_amount / max(payload.monthly_income * 12.0, 1.0)
    cibil_source = "estimated" if cibil_estimated else "provided"

    remarks = [
        (
            f"Default probability is {default_probability:.2f}% from your profile inputs. "
            f"Approval probability is therefore {approval_probability:.2f}% (calculated as 100 - default probability)."
        ),
        (
            f"Key drivers: EMI-to-income is {emi_ratio * 100:.1f}%, expense-to-income is {expense_ratio * 100:.1f}%, "
            f"savings cover is about {savings_months:.1f} month(s), and CIBIL used is {cibil_score_used} ({cibil_source})."
        ),
        condition_msg,
        primary_driver,
        (
            f"Best-fit product for your purpose ({payload.purpose}) is {purpose_loan_type}, "
            f"with suggested tenure around {suggested_tenure} months. "
            f"Requested loan size is {loan_to_annual_income * 100:.1f}% of annual income."
        ),
    ]

    return RiskAssessmentResponse(
        default_probability=default_probability,
        approval_probability=approval_probability,
        risk_category=risk_category,
        cibil_score_used=cibil_score_used,
        cibil_estimated=cibil_estimated,
        remarks=remarks,
        recommended_loan_type=purpose_loan_type,
        suggested_tenure_months=suggested_tenure,
        estimated_monthly_emi=round(estimated_emi, 2),
    )


def train_risk_model_from_csv(contents: bytes, target_column: str | None = None) -> dict[str, Any]:
    df = pd.read_csv(io.BytesIO(contents))
    artifact = risk_model_manager.train(df, target_column=target_column)
    return {
        "best_model": artifact.best_model_name,
        "target_column": artifact.target_column,
        "samples": len(df),
        "metrics": artifact.metrics,
        "trained_features": artifact.trained_features,
    }


def _decode_bytes(contents: bytes) -> str:
    for enc in ["utf-8-sig", "utf-8", "latin1"]:
        try:
            return contents.decode(enc)
        except Exception:
            continue
    return contents.decode("utf-8", errors="ignore")


def _normalize_col_name(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(name).strip().lower())


def _to_numeric(series: pd.Series) -> pd.Series:
    cleaned = (
        series.astype(str)
        .str.replace(",", "", regex=False)
        .str.replace(r"[^\d\.\-]", "", regex=True)
        .replace({"": np.nan, "-": np.nan, ".": np.nan, "--": np.nan})
    )
    return pd.to_numeric(cleaned, errors="coerce")


def _extract_text_transactions(text: str) -> pd.DataFrame:
    date_re = re.compile(
        r"(\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b)",
        flags=re.IGNORECASE,
    )
    amount_re = re.compile(r"[-+]?\d[\d,]*\.?\d*")

    rows: list[dict[str, Any]] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if len(line) < 5:
            continue
        numbers = amount_re.findall(line)
        if not numbers:
            continue

        date_match = date_re.search(line)
        parsed_numbers = []
        for n in numbers:
            try:
                parsed_numbers.append(float(n.replace(",", "")))
            except Exception:
                continue
        if not parsed_numbers:
            continue

        amount = abs(parsed_numbers[0])
        balance = abs(parsed_numbers[-1]) if len(parsed_numbers) >= 2 else np.nan
        lower = line.lower()
        if any(k in lower for k in ["credit", "cr", "salary", "deposit", "refund", "received"]):
            tx_type = "credit"
        elif any(k in lower for k in ["debit", "dr", "withdraw", "spent", "purchase", "payment"]):
            tx_type = "debit"
        else:
            tx_type = "debit"

        rows.append(
            {
                "date": date_match.group(0) if date_match else "",
                "description": line,
                "amount": amount,
                "balance": balance,
                "type": tx_type,
            }
        )

    return pd.DataFrame(rows)


def _tables_from_pdf(contents: bytes) -> tuple[pd.DataFrame | None, str]:
    try:
        import pdfplumber
    except Exception:
        return None, "pdfparser_missing"

    tables: list[pd.DataFrame] = []
    all_text: list[str] = []
    with pdfplumber.open(io.BytesIO(contents)) as pdf:
        for page in pdf.pages:
            all_text.append(page.extract_text() or "")
            for table in page.extract_tables() or []:
                cleaned_rows = [
                    [str(c).strip() if c is not None else "" for c in row]
                    for row in table
                    if row and any((c is not None and str(c).strip()) for c in row)
                ]
                if len(cleaned_rows) < 2:
                    continue
                header = cleaned_rows[0]
                if len(set(header)) == len(header):
                    df = pd.DataFrame(cleaned_rows[1:], columns=header)
                else:
                    df = pd.DataFrame(cleaned_rows)
                tables.append(df)

    if tables:
        return pd.concat(tables, ignore_index=True), "pdf_table"

    text_df = _extract_text_transactions("\n".join(all_text))
    if not text_df.empty:
        return text_df, "pdf_text"
    return None, "pdf_unparsed"


def _tables_from_docx(contents: bytes) -> tuple[pd.DataFrame | None, str]:
    try:
        import docx
    except Exception:
        return None, "docxparser_missing"

    document = docx.Document(io.BytesIO(contents))
    tables: list[pd.DataFrame] = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if len(rows) < 2:
            continue
        header = rows[0]
        if len(set(header)) == len(header):
            df = pd.DataFrame(rows[1:], columns=header)
        else:
            df = pd.DataFrame(rows)
        tables.append(df)

    if tables:
        return pd.concat(tables, ignore_index=True), "docx_table"

    text = "\n".join(p.text for p in document.paragraphs)
    text_df = _extract_text_transactions(text)
    if not text_df.empty:
        return text_df, "docx_text"
    return None, "docx_unparsed"


def _load_statement_dataframe(contents: bytes, filename: str | None = None) -> tuple[pd.DataFrame, str]:
    extension = ""
    if filename and "." in filename:
        extension = filename.lower().rsplit(".", 1)[-1]

    buffer = io.BytesIO(contents)
    text = _decode_bytes(contents)

    def _csv_auto() -> pd.DataFrame:
        return pd.read_csv(io.BytesIO(contents), sep=None, engine="python")

    loaders: list[tuple[str, Any]] = []
    if extension == "csv":
        loaders = [("csv", _csv_auto)]
    elif extension in {"xlsx", "xls"}:
        loaders = [("excel", lambda: pd.read_excel(buffer))]
    elif extension in {"tsv", "txt"}:
        loaders = [("text_csv", _csv_auto)]
    elif extension == "json":
        loaders = [("json", lambda: pd.read_json(io.BytesIO(contents)))]
    elif extension == "xml":
        loaders = [("xml", lambda: pd.read_xml(io.BytesIO(contents)))]
    elif extension == "pdf":
        loaders = [("pdf", lambda: _tables_from_pdf(contents)[0])]
    elif extension in {"docx", "doc"}:
        loaders = [("docx", lambda: _tables_from_docx(contents)[0])]

    # Generic fallback chain for unknown/failed formats.
    loaders.extend(
        [
            ("csv", _csv_auto),
            ("excel", lambda: pd.read_excel(io.BytesIO(contents))),
            ("json", lambda: pd.read_json(io.BytesIO(contents))),
            ("xml", lambda: pd.read_xml(io.BytesIO(contents))),
            ("pdf", lambda: _tables_from_pdf(contents)[0]),
            ("docx", lambda: _tables_from_docx(contents)[0]),
        ]
    )

    for loader_name, fn in loaders:
        try:
            candidate = fn()
            if candidate is not None and isinstance(candidate, pd.DataFrame) and not candidate.empty:
                return candidate, loader_name
        except Exception:
            continue

    text_df = _extract_text_transactions(text)
    if not text_df.empty:
        return text_df, "text_fallback"

    return pd.DataFrame(), "unparsed"


def _best_numeric_column(df: pd.DataFrame, excluded: set[str]) -> str | None:
    best_col = None
    best_score = -1.0
    for col in df.columns:
        if col in excluded:
            continue
        numeric = _to_numeric(df[col])
        score = float(numeric.notna().mean())
        if score > best_score:
            best_score = score
            best_col = col
    if best_score < 0.3:
        return None
    return best_col


def _find_column(df: pd.DataFrame, keywords: list[str]) -> str | None:
    normalized = {col: _normalize_col_name(col) for col in df.columns}
    for key in keywords:
        for col, norm in normalized.items():
            if key in norm:
                return col
    return None


def _normalize_statement(df: pd.DataFrame) -> tuple[pd.DataFrame, dict[str, str | None]]:
    work = df.copy()
    work.columns = [str(c).strip() for c in work.columns]

    date_col = _find_column(work, ["date", "valuedate", "postdate", "txndate", "transactiondate"])
    balance_col = _find_column(
        work,
        ["balance", "runningbalance", "availablebalance", "closingbalance", "runningbal", "closingbal", "availbal", "bal"],
    )
    type_col = _find_column(work, ["drcr", "debitcredit", "transtype", "type"])
    description_col = _find_column(
        work, ["description", "narration", "remarks", "details", "particular", "merchant", "note"]
    )
    debit_col = _find_column(work, ["debit", "withdraw", "dramount", "dr"])
    credit_col = _find_column(work, ["credit", "deposit", "cramount", "cr"])
    amount_col = _find_column(work, ["amount", "txnamount", "transactionamount", "amt"])

    normalized = pd.DataFrame(index=work.index)
    if date_col:
        parsed_dates = pd.to_datetime(work[date_col], errors="coerce", dayfirst=True)
        if parsed_dates.notna().mean() < 0.3:
            parsed_dates = pd.to_datetime(work[date_col], errors="coerce", dayfirst=False)
        normalized["date"] = parsed_dates
    else:
        normalized["date"] = pd.NaT

    if debit_col and credit_col:
        debit = _to_numeric(work[debit_col]).fillna(0).abs()
        credit = _to_numeric(work[credit_col]).fillna(0).abs()
        amount = credit.where(credit > 0, debit)
        tx_type = np.where(credit > 0, "credit", np.where(debit > 0, "debit", "debit"))
        normalized["amount"] = amount
        normalized["type"] = tx_type
    else:
        if amount_col is None:
            amount_col = _best_numeric_column(work, excluded={c for c in [balance_col, date_col] if c})
        amount = _to_numeric(work[amount_col]) if amount_col else pd.Series(np.nan, index=work.index)
        normalized["amount"] = amount.abs()

        if type_col:
            type_values = work[type_col].astype(str).str.strip().str.lower()
            normalized["type"] = np.where(
                type_values.str.contains("cr|credit|dep|salary|refund", regex=True),
                "credit",
                np.where(
                    type_values.str.contains("dr|debit|with|pay|purchase", regex=True),
                    "debit",
                    "debit",
                ),
            )
        else:
            if description_col:
                desc = work[description_col].astype(str).str.lower()
                normalized["type"] = np.where(
                    desc.str.contains("credit|salary|deposit|refund|received"),
                    "credit",
                    np.where(desc.str.contains("debit|withdraw|payment|purchase|spent"), "debit", "debit"),
                )
            else:
                normalized["type"] = np.where(amount < 0, "debit", "credit")

    if balance_col:
        normalized["balance"] = _to_numeric(work[balance_col])
    else:
        normalized["balance"] = np.nan

    if description_col:
        normalized["description"] = work[description_col].astype(str).str.strip()
    else:
        normalized["description"] = ""

    normalized = normalized.dropna(subset=["amount"])
    normalized["amount"] = normalized["amount"].fillna(0)
    normalized["type"] = normalized["type"].fillna("debit").astype(str).str.lower()
    normalized.loc[~normalized["type"].isin(["credit", "debit"]), "type"] = "debit"
    return normalized, {
        "date_col": date_col,
        "amount_col": amount_col if not (debit_col and credit_col) else None,
        "debit_col": debit_col,
        "credit_col": credit_col,
        "type_col": type_col,
        "balance_col": balance_col,
        "description_col": description_col,
    }


def analyze_bank_statement(contents: bytes, filename: str | None = None) -> dict[str, Any]:
    raw_df, _ = _load_statement_dataframe(contents, filename=filename)
    if raw_df.empty:
        return {
            "monthly_income_estimate": 0.0,
            "monthly_expense_estimate": 0.0,
            "avg_monthly_balance": 0.0,
            "income_volatility_index": 0.0,
            "upi_transaction_frequency": 0,
        }

    df, detected = _normalize_statement(raw_df)
    if df.empty:
        return {
            "monthly_income_estimate": 0.0,
            "monthly_expense_estimate": 0.0,
            "avg_monthly_balance": 0.0,
            "income_volatility_index": 0.0,
            "upi_transaction_frequency": 0,
        }

    credit = df[df["type"] == "credit"]["amount"].abs()
    debit = df[df["type"] == "debit"]["amount"].abs()

    monthly_income = float(credit.sum()) if not credit.empty else 0.0
    monthly_expenses = float(debit.sum()) if not debit.empty else 0.0
    avg_monthly_balance = float(df["balance"].dropna().mean()) if df["balance"].notna().any() else 0.0
    income_volatility = float(credit.std(ddof=0)) if len(credit) > 1 else 0.0

    desc_text = df["description"].fillna("").astype(str).str.lower()
    if detected.get("description_col") is not None:
        upi_frequency = int(desc_text.str.contains("upi|gpay|phonepe|paytm|bhim").sum())
    else:
        upi_frequency = int(df["type"].eq("debit").sum())

    return {
        "monthly_income_estimate": round(monthly_income, 2),
        "monthly_expense_estimate": round(monthly_expenses, 2),
        "avg_monthly_balance": round(avg_monthly_balance, 2),
        "income_volatility_index": round(income_volatility, 4),
        "upi_transaction_frequency": upi_frequency,
    }
